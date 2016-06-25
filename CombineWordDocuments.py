#pip install python-docx
from docx import Document
import os

folderpath = raw_input("Enter the path to the folder: ")

dirs = os.listdir(folderpath)

new_doc = Document()

for fn in dirs:
    if fn.endswith('.docx'):
        file_path = os.path.join(folderpath, fn)
        document = Document(file_path)
        new_doc.add_heading(fn,0)
        for para in document.paragraphs:
            new_doc.add_paragraph(para.text)

new_doc.save('combined.docx')